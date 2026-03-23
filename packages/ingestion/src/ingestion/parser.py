from __future__ import annotations

import re
from pathlib import Path

from core_types.enums import ProjectFileType
from docx import Document
import httpx
from pypdf import PdfReader

from ingestion.models import ParsedDocument, ParsedImageAsset


class IngestionParser:
    def parse_file(self, storage_path: str, file_type: str) -> ParsedDocument:
        path = Path(storage_path)
        suffix = path.suffix.lower()
        normalized_type = file_type or ProjectFileType.AUTO_DETECTED

        if normalized_type == ProjectFileType.URL:
            return self._parse_url(storage_path)

        if normalized_type == ProjectFileType.PDF or suffix == '.pdf':
            return self._parse_pdf(path)

        if normalized_type == ProjectFileType.DOCX or suffix == '.docx':
            return self._parse_docx(path)

        if normalized_type in {ProjectFileType.MARKDOWN, ProjectFileType.TXT, ProjectFileType.AUTO_DETECTED} or suffix in {'.md', '.markdown', '.txt'}:
            return self._parse_markdown_like(path)

        raise ValueError(f'Unsupported file type for phase 2 parser skeleton: {file_type}')

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        if not path.exists():
            raise FileNotFoundError(f'Input file not found: {path}')

        reader = PdfReader(str(path))
        warnings: list[str] = []
        page_chunks: list[str] = []

        for index, page in enumerate(reader.pages):
            extracted_text = page.extract_text() or ''
            cleaned_text = extracted_text.strip()
            if cleaned_text:
                page_chunks.append(cleaned_text)
            else:
                warnings.append(f'page {index + 1} could not be parsed')

        raw_text = '\n\n'.join(page_chunks)
        markdown_pages = [f'## Page {index + 1}\n\n{content}' for index, content in enumerate(page_chunks)]
        markdown = '\n\n'.join(markdown_pages)

        if not page_chunks:
            warnings.append('pdf is empty after parsing')

        return ParsedDocument(
            title=path.stem,
            raw_text=raw_text,
            markdown=markdown,
            page_chunks=page_chunks,
            warnings=warnings,
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        if not path.exists():
            raise FileNotFoundError(f'Input file not found: {path}')

        document = Document(str(path))
        warnings: list[str] = []
        markdown_blocks: list[str] = []
        page_chunks: list[str] = []
        raw_text_blocks: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = (paragraph.style.name or '').lower() if paragraph.style else ''
            markdown_text = text
            if style_name.startswith('heading'):
                level_text = ''.join(character for character in style_name if character.isdigit())
                level = int(level_text) if level_text else 1
                level = max(1, min(level, 6))
                markdown_text = f"{'#' * level} {text}"
            elif style_name.startswith('list bullet') or style_name.startswith('list paragraph'):
                markdown_text = f"- {text}"
            elif style_name.startswith('list number'):
                markdown_text = f"1. {text}"

            markdown_blocks.append(markdown_text)
            page_chunks.append(text)
            raw_text_blocks.append(text)

        images: list[ParsedImageAsset] = []
        image_relationships = getattr(document.part, 'related_parts', {})
        image_index = 1
        for relation_name, related_part in image_relationships.items():
            content_type = getattr(related_part, 'content_type', '') or ''
            if not content_type.startswith('image/'):
                continue
            images.append(
                ParsedImageAsset(
                    asset_id=f'image-{image_index}',
                    title=f'Image {image_index}',
                    description=f'Extracted image reference from DOCX relation {relation_name}',
                )
            )
            markdown_blocks.append(f'![Image {image_index}](docx-image:{relation_name})')
            image_index += 1

        markdown = '\n\n'.join(markdown_blocks)
        raw_text = '\n'.join(raw_text_blocks)
        if not markdown_blocks:
            warnings.append('docx is empty after parsing')

        return ParsedDocument(
            title=path.stem,
            raw_text=raw_text,
            markdown=markdown,
            page_chunks=page_chunks,
            images=images,
            warnings=warnings,
        )

    def _parse_url(self, url: str) -> ParsedDocument:
        response = httpx.get(url, follow_redirects=True, timeout=15.0)
        response.raise_for_status()

        html = response.text
        title_match = re.search(r'<title>(.*?)</title>', html, flags=re.IGNORECASE | re.DOTALL)
        title = self._strip_html(title_match.group(1)) if title_match else url
        text = self._extract_text_from_html(html)
        chunks = [chunk.strip() for chunk in text.split('\n\n') if chunk.strip()]
        warnings: list[str] = []
        if not chunks:
            warnings.append('url content is empty after parsing')

        markdown = f'# {title}\n\n{text}'.strip()
        return ParsedDocument(
            title=title,
            raw_text=text,
            markdown=markdown,
            page_chunks=chunks,
            warnings=warnings,
        )

    def _parse_markdown_like(self, path: Path) -> ParsedDocument:
        if not path.exists():
            raise FileNotFoundError(f'Input file not found: {path}')

        text = path.read_text(encoding='utf-8')
        chunks = [chunk.strip() for chunk in text.split('\n\n') if chunk.strip()]
        warnings: list[str] = []
        if not chunks:
            warnings.append('document is empty after parsing')

        return ParsedDocument(
            title=path.stem,
            raw_text=text,
            markdown=text,
            page_chunks=chunks,
            warnings=warnings,
        )

    def _extract_text_from_html(self, html: str) -> str:
        without_scripts = re.sub(r'<script.*?>.*?</script>', ' ', html, flags=re.IGNORECASE | re.DOTALL)
        without_styles = re.sub(r'<style.*?>.*?</style>', ' ', without_scripts, flags=re.IGNORECASE | re.DOTALL)
        with_breaks = re.sub(r'</(p|div|section|article|li|h1|h2|h3|h4|h5|h6)>', '\n\n', without_styles, flags=re.IGNORECASE)
        text = re.sub(r'<[^>]+>', ' ', with_breaks)
        text = self._strip_html(text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    def _strip_html(self, text: str) -> str:
        text = re.sub(r'&nbsp;', ' ', text, flags=re.IGNORECASE)
        text = re.sub(r'&amp;', '&', text, flags=re.IGNORECASE)
        text = re.sub(r'&lt;', '<', text, flags=re.IGNORECASE)
        text = re.sub(r'&gt;', '>', text, flags=re.IGNORECASE)
        return re.sub(r'\s+', ' ', text).strip()
